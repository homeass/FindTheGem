"""
ProteusP DOCX Documentation Generator
설치 및 사용 방법 문서 생성
"""

import datetime
from pathlib import Path
from typing import Optional


def generate_docs(output_path: Optional[str] = None) -> str:
    """
    Generate ProteusP documentation as a DOCX file.
    Returns the path to the generated file.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("Warning: python-docx not installed. Install with: pip install python-docx")
        # Fallback: generate a markdown file instead
        return _generate_md_fallback(output_path)

    doc = Document()

    # ─── Styles ───
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # ─── Cover Page ───
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ProteusP Tablet Edition")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Obsidian RAG Pipeline for Termux-Android")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run(f"Version 1.0.0 — {datetime.date.today().strftime('%Y-%m-%d')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ─── Table of Contents (manual) ───
    toc_heading = doc.add_heading("목차", level=1)
    toc_items = [
        "1. 개요",
        "2. 시스템 요구사항",
        "3. 설치 방법",
        "4. 설정",
        "5. 사용 방법",
        "6. Streamlit GUI 사용법",
        "7. 텔레그램 알림 설정",
        "8. 문제 해결",
        "9. 아키텍처 개요",
        "10. 라이선스",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ─── Section 1: Overview ───
    doc.add_heading("1. 개요", level=1)
    doc.add_paragraph(
        "ProteusP Tablet Edition은 Lenovo Legion Y705 (또는 다른 Android 태블릿)에서 "
        "Obsidian 노트를 로컬 LLM과 연동하여 RAG(Retrieval-Augmented Generation) 시스템을 "
        "구축하는 솔루션입니다."
    )
    doc.add_paragraph(
        "Termux-Ubuntu 환경에서 실행되며, 모든 데이터는 태블릿 내에 안전하게 보관됩니다. "
        "외부 서버로 데이터가 전송되지 않으므로 보안성이 중요한 개인 노트, 업무 문서 등을 "
        "안전하게 검색하고 질의할 수 있습니다."
    )

    # Key features
    doc.add_heading("주요 기능", level=2)
    features = [
        "Obsidian Vault 실시간 동기화 및 인덱싱",
        "마크다운, 위키링크([[Link]]), 태그(#tag) 완벽 파싱",
        "하이브리드 검색 (의미 검색 + 키워드 검색 + 재정렬)",
        "로컬 LLM (Ollama + llama3.2korean3B) 기반 질의응답",
        "Streamlit GUI를 통한 모니터링 및 제어",
        "텔레그램 봇 알림 및 문서 전송",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_page_break()

    # ─── Section 2: Requirements ───
    doc.add_heading("2. 시스템 요구사항", level=1)

    doc.add_heading("하드웨어 (Lenovo Legion Y705)", level=2)
    hw_items = [
        "CPU: ARM64 (Snapdragon 8+ Gen1 이상 권장)",
        "RAM: 8GB 이상 (12GB/24GB 권장)",
        "저장공간: 64GB 이상 여유 공간",
        "Android 13+ (또는 Termux 지원 버전)",
    ]
    for item in hw_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("소프트웨어", level=2)
    sw_items = [
        "Termux (F-Droid 버전 권장)",
        "Termux:Ubuntu (proot-distro)",
        "Python 3.10+",
        "Ollama (Linux ARM64)",
        "Obsidian (Android 앱)",
    ]
    for item in sw_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Python 패키지", level=2)
    doc.add_paragraph(
        "필요한 패키지는 install.sh 스크립트가 자동으로 설치합니다. "
        "주요 패키지는 다음과 같습니다:"
    )
    pkg_items = [
        "chromadb — 벡터 데이터베이스",
        "sentence-transformers — 임베딩 모델 (BAAI/bge-m3)",
        "langchain — 문서 처리 프레임워크",
        "ollama — 로컬 LLM 클라이언트",
        "streamlit — GUI 프레임워크",
        "watchdog — 파일 변경 감지",
        "rank-bm25 + flashrank — 검색 및 재정렬",
        "python-docx — 문서 생성",
    ]
    for item in pkg_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ─── Section 3: Installation ───
    doc.add_heading("3. 설치 방법", level=1)

    doc.add_heading("3.1 Termux-Ubuntu 준비", level=2)
    doc.add_paragraph(
        "Termux가 이미 설치되어 있고, Termux:Ubuntu가 준비되어 있다고 가정합니다."
    )
    code = doc.add_paragraph(
        "# 패키지 업데이트\n"
        "pkg update && pkg upgrade\n\n"
        "# proot-distro 설치\n"
        "pkg install proot-distro\n\n"
        "# Ubuntu 설치\n"
        "proot-distro install ubuntu\n\n"
        "# Ubuntu 접속\n"
        "proot-distro login ubuntu"
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_heading("3.2 Ollama 설치", level=2)
    doc.add_paragraph(
        "Ollama는 로컬 LLM 실행을 위한 필수 구성 요소입니다."
    )
    code = doc.add_paragraph(
        "# ARM64 Linux용 Ollama 설치\n"
        "curl -fsSL https://ollama.com/install.sh | sh\n\n"
        "# 한국어 llama3.2 모델 다운로드\n"
        "ollama pull timHan/llama3.2korean3B\n\n"
        "# Ollama 서버 실행 (백그라운드)\n"
        "ollama serve &"
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_heading("3.3 ProteusP 설치", level=2)
    doc.add_paragraph(
        "ProteusP 패키지를 다운로드하고 설치합니다."
    )
    code = doc.add_paragraph(
        "# ProteusP 디렉토리로 이동\n"
        "cd /path/to/proteusp\n\n"
        "# 설치 스크립트 실행\n"
        "chmod +x install.sh\n"
        "./install.sh\n\n"
        "# 또는 pip로 직접 설치\n"
        "pip install -r requirements.txt\n"
        "pip install -e ."
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_page_break()

    # ─── Section 4: Configuration ───
    doc.add_heading("4. 설정", level=1)

    doc.add_heading("4.1 환경 변수", level=2)
    doc.add_paragraph(
        "주요 설정은 환경 변수로 제어할 수 있습니다:"
    )
    env_items = [
        "PROTEUSP_VAULT_PATH — Obsidian Vault 경로 (기본: ~/storage/shared/Obsidian/Vault)",
        "PROTEUSP_DB_PATH — ChromaDB 저장 경로 (기본: ~/.proteusp/chroma_db)",
        "OLLAMA_HOST — Ollama 서버 주소 (기본: http://localhost:11434)",
        "PROTEUSP_LLM_MODEL — LLM 모델명 (기본: timHan/llama3.2korean3B)",
        "PROTEUSP_CHUNK_SIZE — 청크 크기 (기본: 512)",
        "PROTEUSP_CHUNK_OVERLAP — 청크 중복 크기 (기본: 64)",
    ]
    for item in env_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4.2 설정 파일", level=2)
    doc.add_paragraph(
        "설정은 ~/.proteusp/config.json 파일에 JSON 형식으로 저장됩니다. "
        "GUI의 설정 페이지에서도 수정할 수 있습니다."
    )

    doc.add_page_break()

    # ─── Section 5: Usage ───
    doc.add_heading("5. 사용 방법", level=1)

    doc.add_heading("5.1 빠른 시작", level=2)
    doc.add_paragraph("Streamlit GUI를 실행하여 모든 기능을 사용할 수 있습니다:")
    code = doc.add_paragraph(
        "cd /path/to/proteusp\n"
        "streamlit run proteusp/gui_app.py"
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_heading("5.2 CLI 사용", level=2)
    doc.add_paragraph("GUI 없이 직접 파이프라인을 실행할 수 있습니다:")
    code = doc.add_paragraph(
        "python3 -c \"\n"
        "from proteusp.pipeline import get_pipeline\n"
        "p = get_pipeline()\n"
        "result = p.index_vault()\n"
        "print(result)\n"
        "\"\n\n"
        "# 검색하기\n"
        "python3 -c \"\n"
        "from proteusp.pipeline import get_pipeline\n"
        "p = get_pipeline()\n"
        "r = p.query('프로젝트 일정이 어떻게 되나요?')\n"
        "print(r['response'])\n"
        "\""
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    doc.add_heading("5.3 파일 감시 모드", level=2)
    doc.add_paragraph(
        "Watchdog이 Obsidian Vault의 파일 변경을 감지하여 자동으로 인덱스를 업데이트합니다. "
        "GUI에서 '파일 감시 시작' 버튼을 누르면 활성화됩니다."
    )

    doc.add_page_break()

    # ─── Section 6: Streamlit GUI ───
    doc.add_heading("6. Streamlit GUI 사용법", level=1)

    doc.add_heading("6.1 실행", level=2)
    code = doc.add_paragraph(
        "streamlit run proteusp/gui_app.py --server.port 8501"
    )
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    doc.add_paragraph("브라우저에서 http://localhost:8501 로 접속합니다. "
                       "태블릿 내부에서 접속하거나, 네트워크 설정에 따라 외부에서도 접속 가능합니다.")

    doc.add_heading("6.2 탭 구성", level=2)
    tabs = [
        ("🏠 대시보드", "파이프라인 상태, 인덱스 통계, Ollama 연결 상태를 한눈에 확인"),
        ("🔍 검색", "Obsidian 노트를 검색하고 LLM으로 질의"),
        ("⚙️ 설정", "ProteusP 설정 변경 (Vault 경로, LLM 모델, 청크 크기 등)"),
        ("📡 파일 감시", "Watchdog 파일 감시 상태 및 Git 동기화 설정"),
        ("📨 텔레그램", "텔레그램 봇 설정 및 메시지 전송"),
    ]
    for tab_name, tab_desc in tabs:
        p = doc.add_paragraph()
        run = p.add_run(f"• {tab_name}: ")
        run.bold = True
        p.add_run(tab_desc)

    doc.add_heading("6.3 텔레그램 설정 (첫 실행 시)", level=2)
    steps = [
        "BotFather에서 봇을 생성하고 토큰을 받습니다.",
        "텔레그램 탭에서 봇 토큰을 입력합니다.",
        "시작할 채팅방에서 봇에게 /start를 보냅니다.",
        "GUI에서 '채팅ID 가져오기' 버튼을 누릅니다.",
        "저장하면 메시지와 문서를 텔레그램으로 받을 수 있습니다.",
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_page_break()

    # ─── Section 7: Telegram ───
    doc.add_heading("7. 텔레그램 알림 설정", level=1)
    doc.add_paragraph(
        "ProteusP는 텔레그램 봇을 통해 알림을 보내고 문서를 전송할 수 있습니다. "
        "텔레그램 설정은 Streamlit GUI의 '텔레그램' 탭에서 할 수 있으며, "
        "최초 실행 시 설정 팝업이 자동으로 나타납니다."
    )
    doc.add_paragraph(
        "설정 정보는 ~/.codex/telegram-bridge.json 에 안전하게 저장되어 "
        "다음 번 실행 시 자동으로 불러와집니다."
    )

    doc.add_page_break()

    # ─── Section 8: Troubleshooting ───
    doc.add_heading("8. 문제 해결", level=1)

    problems = [
        ("Ollama 연결 실패",
         "ollama serve가 실행 중인지 확인하세요. "
         "기본 포트 11434가 방화벽에서 열려 있는지 확인하세요."),
        ("벡터 DB 오류",
         "chromadb가 올바르게 설치되었는지 확인하세요. "
         "pip install chromadb 로 재설치해보세요."),
        ("한국어 검색 품질 저하",
         "BAAI/bge-m3 모델은 다국어를 지원하지만, "
         "도메인 특화 용어가 많은 경우 청크 크기를 256으로 줄여보세요."),
        ("메모리 부족",
         "embedding_batch_size를 8 또는 4로 줄이세요. "
         "청크 크기도 256으로 줄이는 것이 좋습니다."),
        ("파일 감지 안됨",
         "PROTEUSP_VAULT_PATH가 올바른지 확인하세요. "
         "Termux의 storage 접근 권한이 있는지 확인하세요."),
    ]
    for prob_title, prob_desc in problems:
        p = doc.add_paragraph()
        run = p.add_run(f"{prob_title}: ")
        run.bold = True
        p.add_run(prob_desc)
        doc.add_paragraph()

    doc.add_page_break()

    # ─── Section 9: Architecture ───
    doc.add_heading("9. 아키텍처 개요", level=1)
    doc.add_paragraph(
        "ProteusP는 모듈식 아키텍처로 설계되었으며, 각 단계가 독립적으로 동작합니다."
    )

    arch_text = (
        "[Obsidian Vault] ──(Watchdog/Git Sync)──> [Parser & Chunker]\n"
        "       │                                            │\n"
        "       │                                    (Metadata & Chunks)\n"
        "       │                                            ▼\n"
        "       │                                    [Embedding (BGE-M3)]\n"
        "       │                                            │\n"
        "       │                                            ▼\n"
        "       │                                    [ChromaDB Vector Store]\n"
        "       │                                            │\n"
        "       │              ┌──────────────────────────────┘\n"
        "       │              │\n"
        "       ▼              ▼\n"
        "[User Query] ──> [Hybrid Searcher (BM25+Vector+Rerank)]\n"
        "                          │\n"
        "                          ▼\n"
        "                   [LLM (Ollama)] ──> [Answer]"
    )
    code = doc.add_paragraph(arch_text)
    code.style = doc.styles["Normal"]
    for run in code.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    doc.add_paragraph()
    doc.add_heading("주요 컴포넌트", level=2)
    components = [
        ("config.py", "설정 관리 (환경 변수, JSON 파일)"),
        ("parser.py", "Obsidian 마크다운 파서 (프론트매터, 태그, 위키링크, 헤더)"),
        ("chunker.py", "헤더 기반 스마트 청킹 + 메타데이터 보존"),
        ("embedder.py", "BGE-M3 임베딩 + 디스크 캐싱"),
        ("vectorstore.py", "ChromaDB 영구 저장소"),
        ("searcher.py", "BM25+벡터 하이브리드 검색 + FlashRank 재정렬"),
        ("llm_service.py", "Ollama API 통신 + RAG 프롬프트 엔지니어링"),
        ("ingestion.py", "Watchdog 파일 감시 + Git 동기화"),
        ("pipeline.py", "전체 파이프라인 오케스트레이션"),
        ("gui_app.py", "Streamlit 웹 GUI"),
        ("telegram_bridge.py", "텔레그램 봇 연동"),
    ]
    for comp, desc in components:
        p = doc.add_paragraph()
        run = p.add_run(f"• {comp}: ")
        run.bold = True
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        p.add_run(desc)

    doc.add_page_break()

    # ─── Section 10: License ───
    doc.add_heading("10. 라이선스", level=1)
    doc.add_paragraph(
        "ProteusP Tablet Edition v1.0.0\n"
        "© 2026. All rights reserved.\n\n"
        "본 소프트웨어는 개인 사용 목적으로 제공됩니다. "
        "무단 배포 및 상업적 사용을 금지합니다."
    )

    # ─── Save ───
    output = output_path or str(Path.cwd() / "ProteusP_Tablet_Edition_Guide.docx")
    doc.save(output)
    print(f"✅ Documentation saved: {output}")
    return output


def _generate_md_fallback(output_path: Optional[str] = None) -> str:
    """Generate markdown documentation as fallback."""
    md_content = """# ProteusP Tablet Edition — 사용 가이드

## 1. 개요
ProteusP Tablet Edition은 Lenovo Legion Y705 태블릿에서 Obsidian 노트를 로컬 LLM과 연동하여
RAG(Retrieval-Augmented Generation) 시스템을 구축하는 솔루션입니다.

## 2. 설치 방법

### 2.1 Ollama 설치
```bash
curl -fsSL https://ollama.com/install.sh | sh
ollama pull timHan/llama3.2korean3B
ollama serve &
```

### 2.2 ProteusP 설치
```bash
cd /path/to/proteusp
pip install -r requirements.txt
```

## 3. 실행
```bash
# GUI 실행
streamlit run proteusp/gui_app.py --server.port 8501

# 또는 CLI로 직접 인덱싱
python3 -c "
from proteusp.pipeline import get_pipeline
p = get_pipeline()
result = p.index_vault()
print(result)
"
```

## 4. 텔레그램 설정
Streamlit GUI 실행 후 '텔레그램' 탭에서:
1. BotFather에서 봇 생성 후 토큰 입력
2. 채팅ID 설정
3. 저장

## 5. 아키텍처
[Obsidian Vault] → [Parser & Chunker] → [Embedding (BGE-M3)] → [ChromaDB]
                                                                    ↓
[User Query] → [Hybrid Searcher (BM25+Vector+Rerank)] → [LLM (Ollama)] → [Answer]
"""
    output = output_path or str(Path.cwd() / "ProteusP_Tablet_Edition_Guide.md")
    Path(output).write_text(md_content, encoding="utf-8")
    print(f"✅ Markdown documentation saved: {output}")
    return output


if __name__ == "__main__":
    generate_docs()
